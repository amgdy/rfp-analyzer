"""Tests for services.utils module."""

import json
import pytest

from services.utils import parse_json_response, format_duration, clean_extracted_text, clean_extracted_markdown


# ============================================================================
# parse_json_response tests
# ============================================================================

class TestParseJsonResponse:
    """Tests for the parse_json_response helper."""

    def test_plain_json(self):
        """Parse plain JSON without any code-block wrappers."""
        data = parse_json_response('{"key": "value"}')
        assert data == {"key": "value"}

    def test_json_with_json_code_block(self):
        """Strip ```json ... ``` wrappers."""
        raw = '```json\n{"a": 1}\n```'
        assert parse_json_response(raw) == {"a": 1}

    def test_json_with_generic_code_block(self):
        """Strip ``` ... ``` wrappers (no language hint)."""
        raw = '```\n{"b": 2}\n```'
        assert parse_json_response(raw) == {"b": 2}

    def test_whitespace_padded(self):
        """Handle leading/trailing whitespace."""
        raw = '  \n ```json\n{"c": 3}\n```  \n'
        assert parse_json_response(raw) == {"c": 3}

    def test_nested_json(self):
        """Parse deeply nested JSON correctly."""
        nested = {"level1": {"level2": {"level3": [1, 2, 3]}}}
        raw = f"```json\n{json.dumps(nested)}\n```"
        assert parse_json_response(raw) == nested

    def test_invalid_json_raises(self):
        """Raise JSONDecodeError on malformed content."""
        with pytest.raises(json.JSONDecodeError):
            parse_json_response("not json at all")

    def test_empty_string_raises(self):
        """Raise JSONDecodeError on empty string."""
        with pytest.raises(json.JSONDecodeError, match="Empty response text"):
            parse_json_response("")

    def test_none_raises(self):
        """Raise JSONDecodeError on None input."""
        with pytest.raises(json.JSONDecodeError, match="Empty response text"):
            parse_json_response(None)

    def test_whitespace_only_raises(self):
        """Raise JSONDecodeError on whitespace-only string."""
        with pytest.raises(json.JSONDecodeError, match="Empty response text"):
            parse_json_response("   \n  ")

    def test_empty_json_object(self):
        """Parse an empty JSON object."""
        assert parse_json_response("{}") == {}

    def test_json_array(self):
        """Parse a JSON array (returns list, which is valid JSON)."""
        result = parse_json_response("[1, 2, 3]")
        assert result == [1, 2, 3]

    def test_only_opening_fence(self):
        """When only opening fence is present, still strips it."""
        # parse_json_response strips startswith ```json but endswith ``` won't match
        # This should still parse if the remaining text is valid JSON
        raw = '```json\n{"d": 4}'
        assert parse_json_response(raw) == {"d": 4}

    def test_real_world_criteria_response(self):
        """Parse a realistic criteria extraction response."""
        response = '''```json
{
    "rfp_title": "Office Renovation RFP",
    "rfp_summary": "Request for proposals for office renovation.",
    "total_weight": 100.0,
    "criteria": [
        {
            "criterion_id": "C-1",
            "name": "Technical Capability",
            "description": "Evaluate technical capabilities",
            "category": "Technical",
            "weight": 40.0,
            "max_score": 100,
            "evaluation_guidance": "Score based on experience"
        },
        {
            "criterion_id": "C-2",
            "name": "Cost",
            "description": "Evaluate pricing",
            "category": "Financial",
            "weight": 60.0,
            "max_score": 100,
            "evaluation_guidance": "Lower is better"
        }
    ],
    "extraction_notes": "Extracted 2 criteria"
}
```'''
        data = parse_json_response(response)
        assert data["rfp_title"] == "Office Renovation RFP"
        assert len(data["criteria"]) == 2
        assert data["criteria"][0]["weight"] == 40.0
        assert data["criteria"][1]["weight"] == 60.0


# ============================================================================
# format_duration tests
# ============================================================================

class TestFormatDuration:
    """Tests for the format_duration helper."""

    def test_zero_seconds(self):
        assert format_duration(0) == "0.0s"

    def test_under_one_second(self):
        assert format_duration(0.5) == "0.5s"

    def test_exactly_one_minute(self):
        assert format_duration(60) == "1m 0.0s"

    def test_one_and_a_half_minutes(self):
        assert format_duration(90) == "1m 30.0s"

    def test_sub_minute(self):
        assert format_duration(45.3) == "45.3s"

    def test_several_minutes(self):
        result = format_duration(125.7)
        assert result == "2m 5.7s"

    def test_large_value(self):
        result = format_duration(3661.2)
        assert result == "61m 1.2s"

    def test_negative_value(self):
        """Negative durations should still format (edge case)."""
        result = format_duration(-5.0)
        assert result == "-5.0s"


# ============================================================================
# clean_extracted_text tests
# ============================================================================

class TestCleanExtractedText:
    """Tests for the clean_extracted_text helper."""

    def test_empty_string(self):
        assert clean_extracted_text("") == ""

    def test_plain_text_unchanged(self):
        text = "This is plain text."
        assert clean_extracted_text(text) == text

    def test_strips_html_tags(self):
        assert clean_extracted_text("<p>Hello</p>") == "Hello"

    def test_strips_nested_html(self):
        result = clean_extracted_text("<div><span>Content</span></div>")
        assert result == "Content"

    def test_strips_html_entities(self):
        result = clean_extracted_text("A&amp;B &lt; C")
        assert "&amp;" not in result
        assert "&lt;" not in result
        assert "A" in result

    def test_strips_numeric_entities(self):
        result = clean_extracted_text("Hello&#160;World")
        assert "&#160;" not in result

    def test_strips_xml_processing_instructions(self):
        result = clean_extracted_text('<?xml version="1.0"?>Hello')
        assert "<?xml" not in result
        assert "Hello" in result

    def test_strips_noisy_image_refs(self):
        result = clean_extracted_text("Before ![image](http://example.com/img.png) After")
        assert "![image]" not in result
        assert "Before" in result
        assert "After" in result

    def test_strips_base64_data_uris(self):
        result = clean_extracted_text("Image: data:image/png;base64,iVBOR... end")
        assert "data:image" not in result

    def test_collapses_excessive_whitespace(self):
        result = clean_extracted_text("word1    word2\t\tword3")
        assert "word1 word2 word3" == result

    def test_collapses_excessive_newlines(self):
        result = clean_extracted_text("para1\n\n\n\n\npara2")
        assert result == "para1\n\npara2"

    def test_strips_line_whitespace(self):
        result = clean_extracted_text("  leading   \n   trailing  ")
        assert result == "leading\ntrailing"

    def test_combined_cleanup(self):
        messy = "<p>Hello &amp; welcome</p>\n\n\n\n<br/>Goodbye"
        result = clean_extracted_text(messy)
        assert "<p>" not in result
        assert "<br/>" not in result
        assert "&amp;" not in result
        assert "Hello" in result
        assert "Goodbye" in result


# ============================================================================
# clean_extracted_markdown tests
# ============================================================================


class TestCleanExtractedMarkdown:
    """Tests for clean_extracted_markdown utility."""

    def test_empty_string(self):
        assert clean_extracted_markdown("") == ""

    def test_none_returns_empty(self):
        assert clean_extracted_markdown(None) == ""

    def test_collapses_excessive_blank_lines(self):
        text = "Line 1\n\n\n\n\n\nLine 2"
        result = clean_extracted_markdown(text)
        assert "\n\n\n\n" not in result
        assert "Line 1" in result and "Line 2" in result

    def test_ensures_blank_line_before_header(self):
        text = "Some paragraph text\n# Header"
        result = clean_extracted_markdown(text)
        assert "\n\n# Header" in result

    def test_normalises_non_breaking_spaces(self):
        text = "Hello\u00a0World"
        result = clean_extracted_markdown(text)
        assert "\u00a0" not in result
        assert "Hello World" in result

    def test_strips_trailing_whitespace(self):
        text = "Line 1   \nLine 2  "
        result = clean_extracted_markdown(text)
        for line in result.split("\n"):
            assert line == line.rstrip()

    def test_preserves_markdown_formatting(self):
        text = "# Header\n\n**bold** and *italic*\n\n| col1 | col2 |\n|---|---|\n| a | b |"
        result = clean_extracted_markdown(text)
        assert "# Header" in result
        assert "**bold**" in result
        assert "| col1 | col2 |" in result


# ============================================================================
# extract_docx_as_markdown tests
# ============================================================================


class TestExtractDocxAsMarkdown:
    """Tests for the extract_docx_as_markdown helper."""

    @staticmethod
    def _make_docx(paragraphs=None, tables=None):
        """Build a minimal DOCX in memory and return its bytes.

        Args:
            paragraphs: list of (text, style_name) tuples.
                        style_name can be "Normal", "Heading 1", etc.
            tables: list of row-lists, e.g. [["H1","H2"],["a","b"]]
        """
        from docx import Document
        import io

        doc = Document()

        if paragraphs:
            for text, style in paragraphs:
                doc.add_paragraph(text, style=style)

        if tables:
            rows = len(tables)
            cols = len(tables[0]) if tables else 0
            tbl = doc.add_table(rows=rows, cols=cols)
            for r_idx, row_data in enumerate(tables):
                for c_idx, cell_text in enumerate(row_data):
                    tbl.rows[r_idx].cells[c_idx].text = cell_text

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_basic_paragraphs(self):
        """Extract plain paragraphs as markdown text."""
        from services.utils import extract_docx_as_markdown

        docx_bytes = self._make_docx(
            paragraphs=[
                ("Hello World", "Normal"),
                ("Second paragraph", "Normal"),
            ]
        )
        result = extract_docx_as_markdown(docx_bytes)
        assert "Hello World" in result
        assert "Second paragraph" in result

    def test_heading_extraction(self):
        """Headings should become markdown headers."""
        from services.utils import extract_docx_as_markdown

        docx_bytes = self._make_docx(
            paragraphs=[
                ("Main Title", "Heading 1"),
                ("Sub Section", "Heading 2"),
                ("Body text", "Normal"),
            ]
        )
        result = extract_docx_as_markdown(docx_bytes)
        assert "# Main Title" in result
        assert "## Sub Section" in result
        assert "Body text" in result

    def test_table_extraction(self):
        """Tables should become markdown tables."""
        from services.utils import extract_docx_as_markdown

        docx_bytes = self._make_docx(
            tables=[
                ["Name", "Score"],
                ["Vendor A", "85"],
                ["Vendor B", "72"],
            ]
        )
        result = extract_docx_as_markdown(docx_bytes)
        assert "| Name | Score |" in result
        assert "| --- | --- |" in result
        assert "| Vendor A | 85 |" in result
        assert "| Vendor B | 72 |" in result

    def test_empty_document(self):
        """An empty DOCX should return an empty string."""
        from services.utils import extract_docx_as_markdown

        docx_bytes = self._make_docx()
        result = extract_docx_as_markdown(docx_bytes)
        assert result == ""

    def test_mixed_content(self):
        """A DOCX with headings, paragraphs, and tables."""
        from services.utils import extract_docx_as_markdown

        docx_bytes = self._make_docx(
            paragraphs=[
                ("Report Title", "Heading 1"),
                ("This is the intro.", "Normal"),
            ],
            tables=[
                ["Criterion", "Weight"],
                ["Cost", "30%"],
            ],
        )
        result = extract_docx_as_markdown(docx_bytes)
        assert "# Report Title" in result
        assert "This is the intro." in result
        assert "| Criterion | Weight |" in result
        assert "| Cost | 30% |" in result

    def test_invalid_bytes_raises_valueerror(self):
        """Non-DOCX bytes should raise a ValueError with a helpful message."""
        from services.utils import extract_docx_as_markdown

        with pytest.raises(ValueError, match="not a valid DOCX document"):
            extract_docx_as_markdown(b"not a docx file")

    def test_old_doc_binary_raises_valueerror(self):
        """An old binary .doc file (non-ZIP) should raise ValueError."""
        from services.utils import extract_docx_as_markdown

        # Old .doc files start with the OLE2 compound file magic bytes
        ole2_header = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 100
        with pytest.raises(ValueError, match="not a valid DOCX"):
            extract_docx_as_markdown(ole2_header)
