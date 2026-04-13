"""
Comparison Agent for Multi-Vendor RFP Evaluation.

This module provides an agent that compares evaluation results
across multiple vendors and generates comparison reports.
"""

import os
import json
import time
import csv
import io
from datetime import datetime
from typing import List, Dict, Any, Optional, Callable

from agent_framework.openai import OpenAIChatClient
from agent_framework import Agent
from azure.identity import DefaultAzureCredential
from pydantic import BaseModel, Field
from dotenv import load_dotenv

from .logging_config import get_logger
from .token_utils import (
    estimate_token_count,
    calculate_token_budget,
    truncate_content,
)
from .utils import parse_json_response
from .retry_utils import run_with_retry, check_for_refusal
from .telemetry import get_tracer

# Optional dependency for Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

load_dotenv()

# Get logger from centralized config
logger = get_logger(__name__)
tracer = get_tracer(__name__)


class VendorRanking(BaseModel):
    """Ranking for a single vendor."""

    rank: int = Field(description="Rank position (1 = best)")
    vendor_name: str = Field(description="Vendor name")
    total_score: float = Field(description="Total weighted score")
    grade: str = Field(description="Letter grade")
    key_strengths: List[str] = Field(description="Top 3 strengths")
    key_concerns: List[str] = Field(description="Top 3 concerns")
    recommendation: str = Field(description="Brief recommendation")


class CriterionComparison(BaseModel):
    """Comparison of vendors for a specific criterion."""

    criterion_id: str = Field(description="Criterion ID")
    criterion_name: str = Field(description="Criterion name")
    weight: float = Field(description="Criterion weight")
    best_vendor: str = Field(description="Best performing vendor for this criterion")
    worst_vendor: str = Field(description="Worst performing vendor for this criterion")
    score_range: str = Field(description="Score range across vendors (e.g., '65-92')")
    insights: str = Field(description="Key insights for this criterion")


class ComparisonResult(BaseModel):
    """Complete comparison result for multiple vendors."""

    rfp_title: str = Field(description="Title of the RFP")
    comparison_date: str = Field(description="Date of comparison")
    total_vendors: int = Field(description="Number of vendors compared")

    # Rankings
    vendor_rankings: List[VendorRanking] = Field(description="Vendors ranked by score")

    # Criterion-level comparison
    criterion_comparisons: List[CriterionComparison] = Field(
        description="Comparison by criterion"
    )

    # Overall analysis
    winner_summary: str = Field(description="Summary of recommended vendor")
    comparison_insights: List[str] = Field(description="Key insights from comparison")
    selection_recommendation: str = Field(description="Final selection recommendation")
    risk_comparison: str = Field(description="Comparative risk assessment")


class ComparisonAgent:
    """
    Agent for comparing multiple vendor evaluations.

    Takes evaluation results from multiple vendors and generates
    comparative analysis, rankings, and recommendations.
    """

    SYSTEM_INSTRUCTIONS = """You are an expert procurement analyst specializing in vendor comparison and selection.

Your task is to analyze evaluation results from multiple vendors responding to the same RFP 
and provide a comprehensive comparative analysis.

## YOUR RESPONSIBILITIES:

1. **Rank Vendors**: Order vendors by total score, identifying the best performer
2. **Compare by Criterion**: Analyze how vendors performed on each evaluation criterion
3. **Identify Patterns**: Find strengths and weaknesses across the vendor pool
4. **Provide Insights**: Offer actionable insights for the selection committee
5. **Make Recommendations**: Provide clear selection recommendations

## ANALYSIS APPROACH:

For each vendor:
- Review their total score and individual criterion scores
- Identify their top 3 strengths and top 3 concerns
- Assess their suitability for the project

For the comparison:
- Identify which vendors excel in which areas
- Note any significant score gaps between vendors
- Highlight criteria where all vendors performed well or poorly
- Consider risk factors and value for money

## OUTPUT FORMAT:

Respond with a valid JSON object:

```json
{
  "rfp_title": "RFP title",
  "comparison_date": "YYYY-MM-DD",
  "total_vendors": <number>,
  "vendor_rankings": [
    {
      "rank": 1,
      "vendor_name": "Vendor Name",
      "total_score": 85.5,
      "grade": "B",
      "key_strengths": ["strength 1", "strength 2", "strength 3"],
      "key_concerns": ["concern 1", "concern 2", "concern 3"],
      "recommendation": "Brief recommendation for this vendor"
    }
  ],
  "criterion_comparisons": [
    {
      "criterion_id": "C-1",
      "criterion_name": "Criterion Name",
      "weight": 20.0,
      "best_vendor": "Best Vendor",
      "worst_vendor": "Worst Vendor",
      "score_range": "65-92",
      "insights": "Key insight for this criterion"
    }
  ],
  "winner_summary": "Summary of why the top vendor is recommended",
  "comparison_insights": [
    "Key insight 1",
    "Key insight 2",
    "Key insight 3"
  ],
  "selection_recommendation": "Clear final recommendation with justification",
  "risk_comparison": "Comparative risk assessment across vendors"
}
```

## IMPORTANT:
- Rank ALL vendors by score
- Compare ALL criteria
- Be objective and fair
- Support recommendations with evidence
- Respond with ONLY valid JSON"""

    def __init__(self):
        """Initialize the comparison agent."""
        logger.info("Initializing ComparisonAgent...")

        self._validate_config()

        endpoint = os.getenv("AZURE_OPENAI_ENDPOINT", "").rstrip("/")
        deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

        self.client = OpenAIChatClient(
            credential=DefaultAzureCredential(),
            azure_endpoint=endpoint,
            model=deployment_name,
            api_version="v1",
        )

        self.deployment_name = deployment_name
        logger.info("ComparisonAgent initialized with endpoint: %s", endpoint)

    def _validate_config(self):
        """Validate required configuration."""
        required_vars = ["AZURE_OPENAI_ENDPOINT", "AZURE_OPENAI_DEPLOYMENT_NAME"]
        missing = [var for var in required_vars if not os.getenv(var)]
        if missing:
            raise ValueError(
                f"Missing required environment variables: {', '.join(missing)}"
            )

    async def compare_evaluations(
        self,
        evaluations: List[Dict[str, Any]],
        rfp_title: str,
        progress_callback: Optional[Callable[[str], None]] = None,
        reasoning_effort: str = "high",
    ) -> Dict[str, Any]:
        """
        Compare multiple vendor evaluations.

        If the formatted evaluations exceed the model context window,
        automatically truncates the summary to fit.

        Args:
            evaluations: List of evaluation results from individual vendor scoring
            rfp_title: Title of the RFP
            progress_callback: Optional callback for progress updates
            reasoning_effort: Reasoning effort level ("low", "medium", "high")

        Returns:
            Dictionary containing comparison results
        """
        start_time = time.time()
        logger.info(
            "Starting comparison of %d vendor evaluations (effort: %s)...",
            len(evaluations),
            reasoning_effort,
        )

        with tracer.start_as_current_span("ComparisonAgent.compare_evaluations") as span:
            span.set_attribute("comparison.vendor_count", len(evaluations))
            span.set_attribute("comparison.reasoning_effort", reasoning_effort)
            span.set_attribute("comparison.rfp_title", rfp_title)

            if progress_callback:
                progress_callback("Preparing vendor comparison...")

            # Format evaluations for the prompt
            evaluations_summary = self._format_evaluations_for_prompt(evaluations)

            # Check token budget and truncate if needed
            budget = calculate_token_budget(self.SYSTEM_INSTRUCTIONS)
            prompt_overhead = estimate_token_count(
                f"Please compare the following vendor evaluations.\n\n"
                f"## RFP TITLE: {rfp_title}\n\n## VENDOR EVALUATIONS:\n\n---\n\n"
                f"REQUIREMENTS:\n1-5...\nRespond with ONLY valid JSON."
            )
            content_budget = budget - prompt_overhead
            summary_tokens = estimate_token_count(evaluations_summary)

            if summary_tokens > content_budget:
                logger.warning(
                    "Comparison summary (~%d tokens) exceeds budget (%d tokens) — truncating",
                    summary_tokens,
                    content_budget,
                )
                evaluations_summary = truncate_content(evaluations_summary, content_budget)

            user_prompt = f"""Please compare the following vendor evaluations and provide a comprehensive analysis.

## RFP TITLE: {rfp_title}

## VENDOR EVALUATIONS:

{evaluations_summary}

---

REQUIREMENTS:
1. Rank all vendors by total score
2. Compare performance on each criterion
3. Identify the best and worst performers per criterion
4. Provide clear selection recommendations
5. Assess comparative risks

Respond with ONLY valid JSON matching the schema in your instructions."""

            try:
                agent = Agent(
                    client=self.client,
                    instructions=self.SYSTEM_INSTRUCTIONS,
                    name="Comparison Agent",
                    default_options={
                        "reasoning": {"effort": reasoning_effort, "summary": "detailed"}
                    },
                )

                if progress_callback:
                    progress_callback("Analyzing vendor comparisons...")

                result = await run_with_retry(
                    lambda: agent.run(user_prompt),
                    description="Vendor comparison",
                )
                response_text = result.text

                # Log token usage
                usage = result.usage_details
                if usage:
                    input_tokens = getattr(usage, "input_token_count", 0) or 0
                    output_tokens = getattr(usage, "output_token_count", 0) or 0
                    total_tokens = getattr(usage, "total_token_count", 0) or 0
                    logger.info(
                        "Comparison - Tokens: Input=%d, Output=%d, Total=%d",
                        input_tokens,
                        output_tokens,
                        total_tokens,
                    )
                    span.set_attribute("llm.input_tokens", input_tokens)
                    span.set_attribute("llm.output_tokens", output_tokens)

                # Parse the response
                comparison_data = self._parse_response(response_text)

                duration = time.time() - start_time
                span.set_attribute("comparison.duration_seconds", duration)
                logger.info("Comparison completed in %.2fs", duration)

                # Add metadata
                comparison_data["_metadata"] = {
                    "comparison_timestamp": datetime.now().isoformat(),
                    "total_duration_seconds": round(duration, 2),
                    "vendors_compared": len(evaluations),
                    "model_deployment": self.deployment_name,
                    "reasoning_effort": reasoning_effort,
                }

                return comparison_data

            except Exception as e:
                span.record_exception(e)
                logger.error("Comparison failed: %s", str(e))
                raise

    def _format_evaluations_for_prompt(self, evaluations: List[Dict[str, Any]]) -> str:
        """Format evaluations for the comparison prompt."""
        formatted = []

        for i, eval_result in enumerate(evaluations, 1):
            vendor_name = eval_result.get("supplier_name", f"Vendor {i}")
            total_score = eval_result.get("total_score", 0)
            grade = eval_result.get("grade", "N/A")

            parts = [f"""### Vendor {i}: {vendor_name}
- **Total Score:** {total_score:.2f}
- **Grade:** {grade}

**Criterion Scores:**
"""]

            # Add criterion scores
            criterion_scores = eval_result.get("criterion_scores", [])
            for cs in criterion_scores:
                parts.append(f"- {cs.get('criterion_name', cs.get('criterion_id', 'Unknown'))}: {cs.get('raw_score', 0):.1f} (weighted: {cs.get('weighted_score', 0):.2f})\n")

            # Add strengths and weaknesses
            strengths = eval_result.get("overall_strengths", [])
            weaknesses = eval_result.get("overall_weaknesses", [])

            if strengths:
                parts.append("\n**Strengths:** " + ", ".join(strengths[:5]))
            if weaknesses:
                parts.append("\n**Weaknesses:** " + ", ".join(weaknesses[:5]))

            formatted.append("".join(parts))

        return "\n\n---\n\n".join(formatted)

    def _parse_response(self, response_text: str) -> Dict[str, Any]:
        """Parse the agent response."""
        # Raise on empty/None so the caller's retry logic can kick in
        if not response_text or not response_text.strip():
            raise RuntimeError(
                "Model returned empty response text — cannot compare vendors"
            )

        # Raise on model refusal so the retry logic retries automatically
        check_for_refusal(response_text)

        try:
            return parse_json_response(response_text)
        except json.JSONDecodeError as e:
            logger.error("Failed to parse comparison JSON: %s", str(e))
            return {
                "rfp_title": "Unknown RFP",
                "comparison_date": datetime.now().strftime("%Y-%m-%d"),
                "total_vendors": 0,
                "vendor_rankings": [],
                "criterion_comparisons": [],
                "winner_summary": "Error parsing comparison results",
                "comparison_insights": [],
                "selection_recommendation": "Unable to provide recommendation due to parsing error",
                "risk_comparison": "Unable to assess",
            }

    def generate_csv_report(
        self, comparison: Dict[str, Any], evaluations: List[Dict[str, Any]]
    ) -> str:
        """
        Generate a CSV report from comparison results.

        Args:
            comparison: Comparison result from compare_evaluations
            evaluations: Original evaluation results

        Returns:
            CSV content as string
        """
        output = io.StringIO()

        # Create comprehensive CSV with all metrics

        # Section 1: Summary Rankings
        writer = csv.writer(output)
        writer.writerow(["RFP Comparison Report"])
        writer.writerow(["RFP Title", comparison.get("rfp_title", "")])
        writer.writerow(["Comparison Date", comparison.get("comparison_date", "")])
        writer.writerow(
            ["Total Vendors", comparison.get("total_vendors", len(evaluations))]
        )
        writer.writerow([])

        # Vendor Rankings
        writer.writerow(["=== VENDOR RANKINGS ==="])
        writer.writerow(
            ["Rank", "Vendor Name", "Total Score", "Grade", "Recommendation"]
        )

        for ranking in comparison.get("vendor_rankings", []):
            writer.writerow(
                [
                    ranking.get("rank", ""),
                    ranking.get("vendor_name", ""),
                    f"{ranking.get('total_score', 0):.2f}",
                    ranking.get("grade", ""),
                    ranking.get("recommendation", ""),
                ]
            )

        writer.writerow([])

        # Criterion Comparison Matrix
        writer.writerow(["=== CRITERION COMPARISON ==="])

        # Get all criteria from first evaluation
        all_criteria = []
        if evaluations:
            all_criteria = [
                cs.get("criterion_name", cs.get("criterion_id", "Unknown"))
                for cs in evaluations[0].get("criterion_scores", [])
            ]

        # Header row
        header = ["Criterion", "Weight"] + [
            e.get("supplier_name", f"Vendor {i + 1}") for i, e in enumerate(evaluations)
        ]
        writer.writerow(header)

        # Score rows for each criterion
        for criterion_idx, criterion_name in enumerate(all_criteria):
            row = [criterion_name]

            # Get weight from first evaluation
            if evaluations and criterion_idx < len(
                evaluations[0].get("criterion_scores", [])
            ):
                weight = evaluations[0]["criterion_scores"][criterion_idx].get(
                    "weight", 0
                )
                row.append(f"{weight:.1f}%")
            else:
                row.append("")

            # Add each vendor's score for this criterion
            for eval_result in evaluations:
                scores = eval_result.get("criterion_scores", [])
                if criterion_idx < len(scores):
                    score = scores[criterion_idx].get("raw_score", 0)
                    row.append(f"{score:.1f}")
                else:
                    row.append("")

            writer.writerow(row)

        # Total scores row
        total_row = ["TOTAL SCORE", "100%"]
        for eval_result in evaluations:
            total_row.append(f"{eval_result.get('total_score', 0):.2f}")
        writer.writerow(total_row)

        writer.writerow([])

        # Insights
        writer.writerow(["=== KEY INSIGHTS ==="])
        for insight in comparison.get("comparison_insights", []):
            writer.writerow([insight])

        writer.writerow([])
        writer.writerow(["=== SELECTION RECOMMENDATION ==="])
        writer.writerow([comparison.get("selection_recommendation", "")])

        return output.getvalue()


def generate_word_report(evaluation: Dict[str, Any], rfp_content: str = "") -> bytes:
    """
    Generate a Word document report from evaluation results.

    Args:
        evaluation: Evaluation result dictionary
        rfp_content: Optional RFP content for context

    Returns:
        Word document as bytes
    """
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not installed. Word export not available.")
        return None

    doc = Document()

    # Title
    title = doc.add_heading("RFP Evaluation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary section
    doc.add_heading("Evaluation Summary", level=1)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"

    summary_data = [
        ("RFP Title", evaluation.get("rfp_title", "")),
        ("Vendor Name", evaluation.get("supplier_name", "")),
        ("Total Score", f"{evaluation.get('total_score', 0):.2f}"),
        ("Grade", evaluation.get("grade", "")),
        ("Evaluation Date", evaluation.get("evaluation_date", "")),
        ("Recommendation", evaluation.get("recommendation", "")),
    ]

    for i, (label, value) in enumerate(summary_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)

    doc.add_paragraph()

    # Criterion Scores Overview
    doc.add_heading("Criterion Scores Overview", level=1)

    criterion_scores = evaluation.get("criterion_scores", [])
    if criterion_scores:
        scores_table = doc.add_table(rows=len(criterion_scores) + 1, cols=4)
        scores_table.style = "Table Grid"

        # Header
        headers = ["Criterion", "Weight", "Score", "Weighted Score"]
        header_row = scores_table.rows[0]
        for i, h in enumerate(headers):
            header_row.cells[i].text = h

        # Data rows
        for i, cs in enumerate(criterion_scores, 1):
            row = scores_table.rows[i]
            row.cells[0].text = cs.get("criterion_name", "")
            row.cells[1].text = f"{cs.get('weight', 0):.1f}%"
            row.cells[2].text = f"{cs.get('raw_score', 0):.1f}"
            row.cells[3].text = f"{cs.get('weighted_score', 0):.2f}"

    doc.add_paragraph()

    # Detailed Criterion Analysis with Justifications
    doc.add_heading("Detailed Criterion Analysis", level=1)

    for cs in criterion_scores:
        criterion_name = cs.get("criterion_name", "Unknown")
        score = cs.get("raw_score", 0)
        weight = cs.get("weight", 0)
        justification = cs.get("justification", "")
        strengths = cs.get("strengths", [])
        gaps = cs.get("gaps", [])

        # Criterion header
        doc.add_heading(f"{criterion_name} (Score: {score:.1f}/100)", level=2)
        doc.add_paragraph(f"Weight: {weight:.1f}%")

        # Justification
        if justification:
            doc.add_heading("Score Justification", level=3)
            doc.add_paragraph(justification)

        # Strengths
        if strengths:
            doc.add_heading("Strengths", level=3)
            for s in strengths:
                doc.add_paragraph(f"• {s}", style="List Bullet")

        # Gaps
        if gaps:
            doc.add_heading("Gaps/Areas for Improvement", level=3)
            for g in gaps:
                doc.add_paragraph(f"• {g}", style="List Bullet")

        doc.add_paragraph()

    # Overall Strengths
    doc.add_heading("Overall Key Strengths", level=1)
    for strength in evaluation.get("overall_strengths", []):
        doc.add_paragraph(f"• {strength}", style="List Bullet")

    # Overall Weaknesses
    doc.add_heading("Overall Key Weaknesses", level=1)
    for weakness in evaluation.get("overall_weaknesses", []):
        doc.add_paragraph(f"• {weakness}", style="List Bullet")

    # Recommendations
    doc.add_heading("Recommendations", level=1)
    for rec in evaluation.get("recommendations", []):
        doc.add_paragraph(f"• {rec}", style="List Bullet")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(evaluation.get("executive_summary", ""))

    # Risk Assessment
    doc.add_heading("Risk Assessment", level=1)
    doc.add_paragraph(evaluation.get("risk_assessment", ""))

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_full_analysis_report(
    comparison: Dict[str, Any], evaluations: List[Dict[str, Any]]
) -> bytes:
    """
    Generate a comprehensive Word document with the full analysis report including comparison.

    Args:
        comparison: Comparison result dictionary
        evaluations: List of evaluation result dictionaries

    Returns:
        Word document as bytes
    """
    if not DOCX_AVAILABLE:
        logger.warning("python-docx not installed. Word export not available.")
        return None

    doc = Document()

    # Title
    title = doc.add_heading("RFP Vendor Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Report metadata
    doc.add_paragraph(
        f"Report Date: {comparison.get('comparison_date', datetime.now().strftime('%Y-%m-%d'))}"
    )
    doc.add_paragraph(f"RFP Title: {comparison.get('rfp_title', 'N/A')}")
    doc.add_paragraph(
        f"Vendors Evaluated: {comparison.get('total_vendors', len(evaluations))}"
    )
    doc.add_paragraph()

    # ==========================================
    # SECTION 1: EXECUTIVE SUMMARY
    # ==========================================
    doc.add_heading("Executive Summary", level=1)

    # Selection Recommendation
    recommendation = comparison.get("selection_recommendation", "")
    if recommendation:
        doc.add_heading("Selection Recommendation", level=2)
        doc.add_paragraph(recommendation)

    # Winner Summary
    winner_summary = comparison.get("winner_summary", "")
    if winner_summary:
        doc.add_heading("Winner Summary", level=2)
        doc.add_paragraph(winner_summary)

    # Key Insights
    insights = comparison.get("comparison_insights", [])
    if insights:
        doc.add_heading("Key Insights", level=2)
        for insight in insights:
            doc.add_paragraph(f"• {insight}", style="List Bullet")

    doc.add_page_break()

    # ==========================================
    # SECTION 2: VENDOR RANKINGS
    # ==========================================
    doc.add_heading("Vendor Rankings", level=1)

    rankings = comparison.get("vendor_rankings", [])
    if rankings:
        # Rankings table
        rankings_table = doc.add_table(rows=len(rankings) + 1, cols=4)
        rankings_table.style = "Table Grid"

        headers = ["Rank", "Vendor", "Score", "Grade"]
        header_row = rankings_table.rows[0]
        for i, h in enumerate(headers):
            header_row.cells[i].text = h

        for i, ranking in enumerate(rankings, 1):
            row = rankings_table.rows[i]
            row.cells[0].text = str(ranking.get("rank", i))
            row.cells[1].text = ranking.get("vendor_name", "")
            row.cells[2].text = f"{ranking.get('total_score', 0):.1f}"
            row.cells[3].text = ranking.get("grade", "")

        doc.add_paragraph()

        # Detailed ranking information
        for ranking in rankings:
            vendor_name = ranking.get("vendor_name", "Unknown")
            rank = ranking.get("rank", 0)

            doc.add_heading(f"#{rank} - {vendor_name}", level=2)

            # Key Strengths
            strengths = ranking.get("key_strengths", [])
            if strengths:
                doc.add_heading("Key Strengths", level=3)
                for s in strengths:
                    doc.add_paragraph(f"• {s}", style="List Bullet")

            # Key Concerns
            concerns = ranking.get("key_concerns", [])
            if concerns:
                doc.add_heading("Key Concerns", level=3)
                for c in concerns:
                    doc.add_paragraph(f"• {c}", style="List Bullet")

            # Recommendation
            rec = ranking.get("recommendation", "")
            if rec:
                doc.add_paragraph(f"Recommendation: {rec}")

            doc.add_paragraph()

    doc.add_page_break()

    # ==========================================
    # SECTION 3: CRITERION COMPARISON
    # ==========================================
    doc.add_heading("Performance by Criterion", level=1)

    criterion_comparisons = comparison.get("criterion_comparisons", [])
    if criterion_comparisons:
        for cc in criterion_comparisons:
            criterion_name = cc.get("criterion_name", "Unknown")
            weight = cc.get("weight", 0)

            doc.add_heading(f"{criterion_name} (Weight: {weight:.1f}%)", level=2)

            # Summary table
            cc_table = doc.add_table(rows=3, cols=2)
            cc_table.style = "Table Grid"
            cc_table.rows[0].cells[0].text = "Best Performer"
            cc_table.rows[0].cells[1].text = cc.get("best_vendor", "N/A")
            cc_table.rows[1].cells[0].text = "Lowest Performer"
            cc_table.rows[1].cells[1].text = cc.get("worst_vendor", "N/A")
            cc_table.rows[2].cells[0].text = "Score Range"
            cc_table.rows[2].cells[1].text = cc.get("score_range", "N/A")

            insights = cc.get("insights", "")
            if insights:
                doc.add_paragraph()
                doc.add_paragraph(f"Insights: {insights}")

            doc.add_paragraph()

    # Score comparison table
    doc.add_heading("Score Comparison Matrix", level=2)

    if evaluations:
        # Get all criteria
        all_criteria = []
        if evaluations[0].get("criterion_scores"):
            all_criteria = evaluations[0]["criterion_scores"]

        if all_criteria:
            # Create comparison table
            num_vendors = len(evaluations)
            matrix_table = doc.add_table(
                rows=len(all_criteria) + 2, cols=num_vendors + 2
            )
            matrix_table.style = "Table Grid"

            # Header row
            matrix_table.rows[0].cells[0].text = "Criterion"
            matrix_table.rows[0].cells[1].text = "Weight"
            for i, eval_result in enumerate(evaluations):
                matrix_table.rows[0].cells[i + 2].text = eval_result.get(
                    "supplier_name", f"Vendor {i + 1}"
                )[:15]

            # Data rows
            for row_idx, criterion in enumerate(all_criteria, 1):
                matrix_table.rows[row_idx].cells[0].text = criterion.get(
                    "criterion_name", ""
                )
                matrix_table.rows[row_idx].cells[
                    1
                ].text = f"{criterion.get('weight', 0):.1f}%"

                for col_idx, eval_result in enumerate(evaluations):
                    scores = eval_result.get("criterion_scores", [])
                    if row_idx - 1 < len(scores):
                        score = scores[row_idx - 1].get("raw_score", 0)
                        matrix_table.rows[row_idx].cells[
                            col_idx + 2
                        ].text = f"{score:.1f}"

            # Total row
            last_row = len(all_criteria) + 1
            matrix_table.rows[last_row].cells[0].text = "TOTAL SCORE"
            matrix_table.rows[last_row].cells[1].text = "100%"
            for col_idx, eval_result in enumerate(evaluations):
                matrix_table.rows[last_row].cells[
                    col_idx + 2
                ].text = f"{eval_result.get('total_score', 0):.1f}"

    doc.add_page_break()

    # ==========================================
    # SECTION 4: RISK ASSESSMENT
    # ==========================================
    doc.add_heading("Risk Assessment", level=1)

    risk_comparison = comparison.get("risk_comparison", "")
    if risk_comparison:
        doc.add_paragraph(risk_comparison)
    else:
        doc.add_paragraph("No specific risk assessment provided.")

    doc.add_page_break()

    # ==========================================
    # SECTION 5: DETAILED VENDOR REPORTS
    # ==========================================
    doc.add_heading("Detailed Vendor Reports", level=1)

    for eval_result in evaluations:
        vendor_name = eval_result.get("supplier_name", "Unknown")
        total_score = eval_result.get("total_score", 0)
        grade = eval_result.get("grade", "N/A")

        doc.add_heading(f"{vendor_name}", level=2)

        # Summary
        doc.add_paragraph(f"Total Score: {total_score:.1f} | Grade: {grade}")

        # Executive Summary
        exec_summary = eval_result.get("executive_summary", "")
        if exec_summary:
            doc.add_heading("Executive Summary", level=3)
            doc.add_paragraph(exec_summary)

        # Criterion Scores with Justifications
        doc.add_heading("Criterion Analysis", level=3)

        criterion_scores = eval_result.get("criterion_scores", [])
        for cs in criterion_scores:
            criterion_name = cs.get("criterion_name", "Unknown")
            score = cs.get("raw_score", 0)
            weight = cs.get("weight", 0)
            justification = cs.get("justification", "")
            strengths = cs.get("strengths", [])
            gaps = cs.get("gaps", [])

            doc.add_heading(
                f"{criterion_name}: {score:.1f}/100 (Weight: {weight:.1f}%)", level=4
            )

            if justification:
                doc.add_paragraph(f"Justification: {justification}")

            if strengths:
                doc.add_paragraph("Strengths:")
                for s in strengths:
                    doc.add_paragraph(f"  • {s}")

            if gaps:
                doc.add_paragraph("Gaps:")
                for g in gaps:
                    doc.add_paragraph(f"  • {g}")

        # Overall Strengths
        overall_strengths = eval_result.get("overall_strengths", [])
        if overall_strengths:
            doc.add_heading("Overall Strengths", level=3)
            for s in overall_strengths:
                doc.add_paragraph(f"• {s}", style="List Bullet")

        # Overall Weaknesses
        overall_weaknesses = eval_result.get("overall_weaknesses", [])
        if overall_weaknesses:
            doc.add_heading("Overall Weaknesses", level=3)
            for w in overall_weaknesses:
                doc.add_paragraph(f"• {w}", style="List Bullet")

        # Recommendations
        recommendations = eval_result.get("recommendations", [])
        if recommendations:
            doc.add_heading("Recommendations", level=3)
            for rec in recommendations:
                doc.add_paragraph(f"• {rec}", style="List Bullet")

        doc.add_page_break()

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
