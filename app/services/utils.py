"""Shared utilities for RFP Analyzer services."""

import io
import json
import re
from typing import Any


def parse_json_response(text: str) -> dict[str, Any]:
    """Strip markdown code blocks and parse JSON response text.

    Args:
        text: Raw response text that may contain markdown code blocks

    Returns:
        Parsed JSON as a dictionary

    Raises:
        json.JSONDecodeError: If the text is not valid JSON after stripping
    """
    if not text or not text.strip():
        raise json.JSONDecodeError(
            "Empty response text — the model returned no content",
            doc=text or "",
            pos=0,
        )
    text = text.strip()

    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    text = text.strip()

    return json.loads(text)


def format_duration(seconds: float) -> str:
    """Format duration in minutes and seconds.

    Args:
        seconds: Duration in seconds

    Returns:
        Formatted string like "1m 30.0s" or "45.0s" for durations under a minute
    """
    if seconds < 60:
        return f"{seconds:.1f}s"
    minutes = int(seconds // 60)
    remaining_seconds = seconds % 60
    return f"{minutes}m {remaining_seconds:.1f}s"


def clean_extracted_markdown(markdown: str) -> str:
    """Clean extracted markdown from Azure document processors.

    Fixes common issues produced by Document Intelligence and
    Content Understanding:
    - Collapses excessive blank lines
    - Ensures headers have surrounding blank lines
    - Normalises non-breaking / special whitespace
    - Strips trailing whitespace per line

    Unlike :func:`clean_extracted_text`, this preserves full markdown
    formatting (headers, tables, emphasis, etc.) and is intended for
    the extracted content that will be fed to AI agents.

    Args:
        markdown: Raw markdown from an Azure document processor

    Returns:
        Cleaned markdown string
    """
    if not markdown:
        return ""

    # Collapse runs of 3+ blank lines into 2
    cleaned = re.sub(r"\n{4,}", "\n\n\n", markdown)

    # Ensure headers (#) are preceded by a blank line unless at start
    cleaned = re.sub(r"([^\n])\n(#{1,6} )", r"\1\n\n\2", cleaned)

    # Normalise non-breaking / special spaces
    cleaned = cleaned.replace("\u00a0", " ")

    # Strip trailing whitespace on each line
    cleaned = "\n".join(line.rstrip() for line in cleaned.split("\n"))

    return cleaned.strip()


def clean_extracted_text(text: str) -> str:
    """Clean extracted document text for executive-ready display.

    Strips HTML tags, excessive whitespace, stray markup artifacts,
    and normalises the text so it reads cleanly in a Streamlit UI.

    Args:
        text: Raw extracted text that may contain HTML/markup

    Returns:
        Cleaned plain text suitable for display
    """
    if not text:
        return ""

    # Remove HTML tags
    cleaned = re.sub(r"<[^>]+>", "", text)

    # Remove HTML entities
    cleaned = re.sub(r"&[a-zA-Z]+;", " ", cleaned)
    cleaned = re.sub(r"&#\d+;", " ", cleaned)

    # Remove XML processing instructions / CDATA
    cleaned = re.sub(r"<!\[CDATA\[.*?\]\]>", "", cleaned, flags=re.DOTALL)
    cleaned = re.sub(r"<\?xml[^>]*\?>", "", cleaned)

    # Remove stray markdown image references that are just noise
    cleaned = re.sub(r"!\[(?:image|figure|img)?\]\([^)]*\)", "", cleaned)

    # Remove base64 data URIs (very long inline images)
    cleaned = re.sub(r"data:[a-zA-Z0-9/+;,=]+", "", cleaned)

    # Collapse runs of whitespace (but preserve paragraph breaks)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

    # Strip leading/trailing whitespace from each line
    cleaned = "\n".join(line.strip() for line in cleaned.splitlines())

    # Remove fully blank lines at start/end
    cleaned = cleaned.strip()

    return cleaned


def check_document_protection(file_bytes: bytes, filename: str) -> None:
    """Check whether a document is encrypted or IRM-protected.

    Inspects PDF and DOCX files *before* sending them to an extraction
    service so the user gets a clear, actionable error message instead
    of a cryptic API failure.

    Args:
        file_bytes: Raw document content.
        filename: Original filename (used to determine format).

    Raises:
        ValueError: If the document is encrypted / IRM-protected.
    """
    extension = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    _PROTECTION_MSG = (
        "This document appears to be protected (encrypted or IRM-protected). "
        "Please remove the protection in the original application and upload "
        "the file again."
    )

    if extension == "pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_bytes))
            if reader.is_encrypted:
                raise ValueError(_PROTECTION_MSG)
        except ValueError:
            raise  # re-raise our own ValueError
        except Exception:
            # If pypdf cannot parse the file at all we let the downstream
            # extraction service report the real error.
            pass

    elif extension == "docx":
        try:
            import msoffcrypto

            docx_stream = io.BytesIO(file_bytes)
            office_file = msoffcrypto.OfficeFile(docx_stream)
            if office_file.is_encrypted():
                raise ValueError(_PROTECTION_MSG)
        except ValueError:
            raise  # re-raise our own ValueError
        except Exception:
            # Not a valid OLE/OOXML container — let the downstream
            # handler report the real error.
            pass


def get_pdf_page_count(file_bytes: bytes) -> int:
    """Return the number of pages in a PDF file.

    Args:
        file_bytes: Raw PDF content.

    Returns:
        Page count, or 0 if the file cannot be parsed.
    """
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_bytes))
        return len(reader.pages)
    except Exception:
        return 0


def split_pdf_bytes(file_bytes: bytes, max_pages: int) -> list[bytes]:
    """Split a PDF into chunks of at most *max_pages* pages.

    Each chunk is returned as a standalone PDF (bytes).

    Args:
        file_bytes: The full PDF document as bytes.
        max_pages: Maximum pages per chunk.

    Returns:
        A list of PDF byte-strings, each containing ≤ max_pages pages.
    """
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(io.BytesIO(file_bytes))
    total_pages = len(reader.pages)

    if total_pages <= max_pages:
        return [file_bytes]

    chunks: list[bytes] = []
    for start in range(0, total_pages, max_pages):
        writer = PdfWriter()
        for page_idx in range(start, min(start + max_pages, total_pages)):
            writer.add_page(reader.pages[page_idx])
        buf = io.BytesIO()
        writer.write(buf)
        chunks.append(buf.getvalue())

    return chunks


def extract_docx_as_markdown(file_bytes: bytes) -> str:
    """Extract content from a DOCX file as markdown text.

    Uses ``python-docx`` to read paragraphs and tables and produces a
    markdown representation suitable for downstream AI processing.

    Azure Document Intelligence does not natively support ``.docx`` files,
    so this function is used as a local extraction fallback.

    Args:
        file_bytes: The DOCX document content as bytes

    Returns:
        Extracted content as a markdown string

    Raises:
        ImportError: If ``python-docx`` is not installed
        ValueError: If the file is not a valid DOCX (e.g. old binary
            ``.doc`` format, corrupted archive, or non-ZIP data).
    """
    import zipfile

    from docx import Document  # python-docx

    _MIN_HEADING_LEVEL = 1
    _MAX_HEADING_LEVEL = 6

    try:
        doc = Document(io.BytesIO(file_bytes))
    except zipfile.BadZipFile:
        raise ValueError(
            "The file is not a valid DOCX document. It may be an older "
            "binary .doc format or a corrupted file. Please convert it "
            "to .docx or .pdf and try again."
        )
    except Exception as exc:
        raise ValueError(
            f"Failed to parse DOCX file: {exc}. Please ensure the file "
            "is a valid .docx document or convert it to .pdf."
        ) from exc
    parts: list[str] = []

    # Pre-build element→object maps to avoid O(n²) inner scans
    para_map = {para._element: para for para in doc.paragraphs}
    table_map = {table._element: table for table in doc.tables}

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]  # strip namespace

        if tag == "p":
            para = para_map.get(element)
            if para is None:
                continue
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower() if para.style else ""
            if style_name.startswith("heading"):
                # Map heading level: "Heading 1" → "#", "Heading 2" → "##", etc.
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = _MIN_HEADING_LEVEL
                level = max(_MIN_HEADING_LEVEL, min(level, _MAX_HEADING_LEVEL))
                parts.append(f"\n{'#' * level} {text}\n")
            else:
                parts.append(text)

        elif tag == "tbl":
            table = table_map.get(element)
            if table is not None:
                _render_table(table, parts)

    return clean_extracted_markdown("\n\n".join(parts))


def _render_table(table, parts: list[str]) -> None:
    """Render a python-docx Table as a markdown table into *parts*."""
    rows = table.rows
    if not rows:
        return

    # Build rows of cell text
    md_rows: list[list[str]] = []
    for row in rows:
        md_rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])

    if not md_rows:
        return

    # Header row
    header = "| " + " | ".join(md_rows[0]) + " |"
    separator = "| " + " | ".join("---" for _ in md_rows[0]) + " |"
    body_lines = [
        "| " + " | ".join(r) + " |"
        for r in md_rows[1:]
    ]

    parts.append("\n".join([header, separator, *body_lines]))
